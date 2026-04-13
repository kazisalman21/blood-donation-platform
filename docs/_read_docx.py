import docx
doc = docx.Document(r"E:\coding\blood-donation-platform\docs\CSE470 Project Outline Spring'26 Onwards.docx")
for p in doc.paragraphs:
    print(p.text)
# Also read tables
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i+1} ---")
    for row in table.rows:
        print(" | ".join(cell.text.strip() for cell in row.cells))
